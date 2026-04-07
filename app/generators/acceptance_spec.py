"""Acceptance Specification Generator (Word output).

Generates a 5-section acceptance specification document matching the sample format:
1. 維修設備基本資訊 (table)
2. 施工處置與安全規範 (sub-headed narrative paragraphs)
3. 施工與驗收標準表 (table)
4. 施工保險要求 (paragraph)
5. 施工區域 (paragraph)
"""

import io
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

from app.generators.base import BaseGenerator
from app.llm.batch_processor import LLMBatchProcessor
from app.models.quote_data import AcceptanceContent, QuoteData

logger = logging.getLogger(__name__)

# Blue color used for section headings
HEADING_COLOR = RGBColor(0x4F, 0x81, 0xBD)


class AcceptanceSpecGenerator(BaseGenerator):
    doc_type = "acceptance_spec"
    output_extension = "docx"

    async def generate(self, quote_data: QuoteData, **kwargs) -> bytes:
        """Generate acceptance specification document."""
        acceptance_content = kwargs.get("acceptance_content")
        if not acceptance_content and self.llm:
            acceptance_content = await self.llm.generate_acceptance_content(quote_data)
        if not acceptance_content:
            from app.llm.batch_processor import _fallback_acceptance_content
            acceptance_content = _fallback_acceptance_content(quote_data)

        if self.template_path and self.template_path.exists():
            return self._fill_template(quote_data, acceptance_content)
        else:
            return self._generate_from_scratch(quote_data, acceptance_content)

    def _fill_template(
        self, quote_data: QuoteData, content: AcceptanceContent
    ) -> bytes:
        """Fill an existing Word template with content."""
        from app.templates.engine import WordTemplateEngine

        engine = WordTemplateEngine(self.template_path)
        meta = quote_data.metadata

        engine.replace_placeholder("PROJECT_NAME", meta.project_name)
        engine.replace_placeholder("EQUIPMENT_LOCATION", content.equipment_location)
        engine.replace_placeholder("EQUIPMENT_TYPE", content.equipment_type)
        engine.replace_placeholder("SPEC_INFO", content.spec_info)
        engine.replace_placeholder("PROBLEM_DESCRIPTION", content.problem_description)
        engine.replace_placeholder("REPAIR_NEEDS", content.repair_needs)

        return engine.save_to_bytes()

    def _generate_from_scratch(
        self, quote_data: QuoteData, content: AcceptanceContent
    ) -> bytes:
        """Generate a complete acceptance spec document matching the sample format."""
        doc = Document()

        # Page setup
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

        meta = quote_data.metadata
        doc_num = meta.document_number or ""
        project_display = meta.project_name
        if doc_num:
            project_display = f"（{doc_num}）{project_display}"

        # ========== Title ==========
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"{project_display}驗收規範")
        run.font.size = Pt(16)
        run.bold = True

        # ========== Section 1: 維修設備基本資訊 ==========
        _add_section_heading(doc, "一、維修設備基本資訊")
        _add_info_table(doc, content)

        # ========== Section 2: 施工處置與安全規範 ==========
        _add_section_heading(doc, "二、施工處置與安全規範")

        # Sub-heading: 施工廠商需具備：
        _add_sub_heading(doc, "施工廠商需具備：")
        for text in content.contractor_requirements:
            _add_body_paragraph(doc, text)

        # Sub-heading: 作業前準備：
        _add_sub_heading(doc, "作業前準備：")
        for text in content.preparation_steps:
            _add_body_paragraph(doc, text)

        # Sub-heading: 施工內容：
        _add_sub_heading(doc, "施工內容：")
        for text in content.construction_content:
            _add_body_paragraph(doc, text)

        # Sub-heading: 驗收作業：
        _add_sub_heading(doc, "驗收作業：")
        for text in content.acceptance_procedures:
            _add_body_paragraph(doc, text)

        # ========== Section 3: 施工與驗收標準表 ==========
        _add_section_heading(doc, "三、施工與驗收標準表")
        _add_acceptance_table(doc, content)

        # ========== Section 4: 施工保險要求 ==========
        _add_section_heading(doc, "四、施工保險要求")
        _add_body_paragraph(doc, content.insurance_text)

        # ========== Section 5: 施工區域 ==========
        _add_section_heading(doc, "五、施工區域")
        _add_body_paragraph(doc, content.work_area_text)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def _add_section_heading(doc: Document, text: str):
    """Add a bold section heading (一、二、三...) with blue color."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = HEADING_COLOR
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_sub_heading(doc: Document, text: str):
    """Add a bold sub-heading within a section."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def _add_body_paragraph(doc: Document, text: str):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)


def _add_info_table(doc: Document, content: AcceptanceContent):
    """Add the 維修設備基本資訊 table (5 rows x 2 cols)."""
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.0)

    rows_data = [
        ("設備位置", content.equipment_location),
        ("設備類型", content.equipment_type),
        ("馬力規格", content.spec_info),
        ("問題說明", content.problem_description),
        ("維修需求", content.repair_needs),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        # Label cell
        cell0 = row.cells[0]
        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(label)
        run0.bold = True
        run0.font.size = Pt(11)

        # Value cell
        cell1 = row.cells[1]
        cell1.text = ""
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(value)
        run1.font.size = Pt(11)

    doc.add_paragraph()  # Spacing


def _add_acceptance_table(doc: Document, content: AcceptanceContent):
    """Add the 施工與驗收標準表 (驗收項目 | 驗收標準說明 | 驗收結果 | 備註)."""
    criteria = content.acceptance_criteria
    if not criteria:
        doc.add_paragraph("（無驗收標準資料）")
        return

    num_rows = 1 + len(criteria)
    table = doc.add_table(rows=num_rows, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["驗收項目", "驗收標準說明", "驗收結果", "備註"]
    header_row = table.rows[0]
    for col_idx, header_text in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(11.0)
        row.cells[2].width = Cm(2.0)
        row.cells[3].width = Cm(1.5)

    # Data rows
    for i, criterion in enumerate(criteria):
        row = table.rows[i + 1]
        # Category
        cell0 = row.cells[0]
        cell0.text = ""
        run0 = cell0.paragraphs[0].add_run(criterion.category)
        run0.font.size = Pt(10)

        # Standard description
        cell1 = row.cells[1]
        cell1.text = ""
        run1 = cell1.paragraphs[0].add_run(criterion.standard)
        run1.font.size = Pt(10)

        # Empty result and remark columns
        row.cells[2].text = ""
        row.cells[3].text = ""

    doc.add_paragraph()
