"""Template loading and filling utilities for Word and Excel."""

import copy
import logging
import re
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document
from docx.shared import Pt
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.worksheet import Worksheet

logger = logging.getLogger(__name__)


# ---------- Word Template Engine ----------

class WordTemplateEngine:
    """Fill Word (.docx) templates with structured data."""

    def __init__(self, template_path: Path):
        self.template_path = template_path
        self.doc = Document(str(template_path))

    def replace_placeholder(self, placeholder: str, value: str):
        """Replace a {{placeholder}} in paragraphs and tables."""
        tag = "{{" + placeholder + "}}"
        self._replace_in_paragraphs(self.doc.paragraphs, tag, value)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_paragraphs(cell.paragraphs, tag, value)

    def replace_placeholder_with_list(
        self, placeholder: str, items: list[str], numbered: bool = False
    ):
        """Replace a placeholder with a bulleted/numbered list."""
        tag = "{{" + placeholder + "}}"
        for para in self.doc.paragraphs:
            if tag in para.text:
                # Clear the placeholder paragraph
                para.text = ""
                # Add items as new runs in the same paragraph or new paragraphs
                for i, item in enumerate(items):
                    prefix = f"{i + 1}. " if numbered else "• "
                    if i == 0:
                        run = para.add_run(prefix + item)
                        run.font.size = Pt(11)
                    else:
                        new_para = _insert_paragraph_after(para, prefix + item)
                        new_para.style = para.style
                        para = new_para  # track for next insertion
                return

    def fill_table_rows(
        self,
        table_index: int,
        template_row_index: int,
        data_rows: list[dict[int, str]],
    ):
        """
        Clone a template row in a table and fill with data.

        data_rows: list of {column_index: value} mappings
        """
        if table_index >= len(self.doc.tables):
            logger.warning("Table index %d out of range", table_index)
            return

        table = self.doc.tables[table_index]
        if template_row_index >= len(table.rows):
            logger.warning("Row index %d out of range", template_row_index)
            return

        template_row = table.rows[template_row_index]

        for data in data_rows:
            new_row = _copy_row(table, template_row)
            for col_idx, value in data.items():
                if col_idx < len(new_row.cells):
                    new_row.cells[col_idx].text = str(value)

    def add_paragraph_after_placeholder(
        self, placeholder: str, text: str, bold: bool = False
    ):
        """Add a paragraph after the one containing the placeholder."""
        tag = "{{" + placeholder + "}}"
        for para in self.doc.paragraphs:
            if tag in para.text:
                para.text = para.text.replace(tag, "")
                new_para = _insert_paragraph_after(para, text)
                if bold:
                    for run in new_para.runs:
                        run.bold = True
                return

    def save(self, output_path: Path):
        """Save the filled document."""
        self.doc.save(str(output_path))

    def save_to_bytes(self) -> bytes:
        """Save the document to bytes."""
        import io
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def _replace_in_paragraphs(paragraphs, tag: str, value: str):
        """Replace tag text across runs in paragraphs, preserving formatting."""
        for para in paragraphs:
            if tag in para.text:
                # Try run-by-run replacement first
                full_text = "".join(run.text for run in para.runs)
                if tag in full_text:
                    # Simple case: tag is within runs
                    _replace_across_runs(para.runs, tag, value)
                else:
                    # Fallback: direct text replacement
                    para.text = para.text.replace(tag, value)


# ---------- Excel Template Engine ----------

class ExcelTemplateEngine:
    """Fill Excel (.xlsx) templates with structured data."""

    def __init__(self, template_path: Path):
        self.template_path = template_path
        self.wb = openpyxl.load_workbook(str(template_path))

    def fill_cell(self, sheet_name: str | None, cell_ref: str, value: Any):
        """Fill a specific cell by reference (e.g., 'B3')."""
        ws = self._get_sheet(sheet_name)
        ws[cell_ref] = value

    def fill_cells(self, sheet_name: str | None, mapping: dict[str, Any]):
        """Fill multiple cells from a {cell_ref: value} dict."""
        ws = self._get_sheet(sheet_name)
        for cell_ref, value in mapping.items():
            ws[cell_ref] = value

    def fill_rows(
        self,
        sheet_name: str | None,
        start_row: int,
        columns: dict[str, str],
        data: list[dict[str, Any]],
        preserve_formulas: bool = True,
    ) -> int:
        """
        Fill rows starting from start_row.

        columns: {field_name: column_letter} mapping
        data: list of {field_name: value} dicts
        Returns the next available row after insertion.
        """
        ws = self._get_sheet(sheet_name)

        # If we need to insert rows, copy formatting from the template row
        template_row_styles = _capture_row_styles(ws, start_row)

        for i, record in enumerate(data):
            row_num = start_row + i

            # Insert a new row if needed (beyond the first template row)
            if i > 0 and row_num <= ws.max_row:
                ws.insert_rows(row_num)
                _apply_row_styles(ws, row_num, template_row_styles)

            for field, col_letter in columns.items():
                if field in record:
                    cell = ws[f"{col_letter}{row_num}"]
                    value = record[field]
                    if preserve_formulas and isinstance(
                        cell.value, str
                    ) and cell.value.startswith("="):
                        continue  # Don't overwrite formulas
                    cell.value = value

        return start_row + len(data)

    def clear_column_in_range(
        self, sheet_name: str | None, column: str, start_row: int, end_row: int
    ):
        """Clear values in a column range (e.g., for blanking prices)."""
        ws = self._get_sheet(sheet_name)
        for row in range(start_row, end_row + 1):
            ws[f"{column}{row}"] = None

    def delete_rows(
        self, sheet_name: str | None, row_indices: list[int]
    ):
        """Delete rows by index (in descending order to preserve indices)."""
        ws = self._get_sheet(sheet_name)
        for idx in sorted(row_indices, reverse=True):
            ws.delete_rows(idx)

    def get_all_rows(
        self, sheet_name: str | None, start_row: int = 1
    ) -> list[list[Any]]:
        """Read all rows from a sheet starting from start_row."""
        ws = self._get_sheet(sheet_name)
        rows = []
        for row in ws.iter_rows(min_row=start_row, values_only=True):
            rows.append(list(row))
        return rows

    def get_sheet_names(self) -> list[str]:
        return self.wb.sheetnames

    def save(self, output_path: Path):
        self.wb.save(str(output_path))

    def save_to_bytes(self) -> bytes:
        import io
        buffer = io.BytesIO()
        self.wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_sheet(self, name: str | None) -> Worksheet:
        if name is None:
            return self.wb.active
        return self.wb[name]


# ---------- Helper Functions ----------

def _insert_paragraph_after(paragraph, text: str):
    """Insert a new paragraph after the given paragraph."""
    from docx.oxml.ns import qn

    new_p = copy.deepcopy(paragraph._element)
    new_p.getparent().insert(
        list(new_p.getparent()).index(paragraph._element) + 1, new_p
    )
    # Clear and set text
    for child in new_p.findall(qn("w:r")):
        new_p.remove(child)
    new_run = copy.deepcopy(paragraph.runs[0]._element) if paragraph.runs else None
    if new_run is not None:
        new_run.text = text
        new_p.append(new_run)
    else:
        from docx.oxml import OxmlElement
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        new_p.append(r)

    # Return as Paragraph object
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _replace_across_runs(runs, tag: str, replacement: str):
    """Replace a tag that might span across multiple runs."""
    combined = ""
    run_map = []  # (run_index, start_in_combined, end_in_combined)

    for i, run in enumerate(runs):
        start = len(combined)
        combined += run.text
        run_map.append((i, start, len(combined)))

    idx = combined.find(tag)
    if idx == -1:
        return

    new_combined = combined[:idx] + replacement + combined[idx + len(tag):]

    # Redistribute text across runs
    pos = 0
    for i, run in enumerate(runs):
        _, start, end = run_map[i]
        orig_len = end - start
        if pos < len(new_combined):
            run.text = new_combined[pos:pos + orig_len + (len(new_combined) - len(combined))]
            pos += len(run.text)
        else:
            run.text = ""

    # Simple fallback: put all text in first run, clear others
    if "".join(r.text for r in runs) != new_combined:
        runs[0].text = new_combined
        for r in runs[1:]:
            r.text = ""


def _copy_row(table, source_row):
    """Copy a table row and append to the table."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    new_tr = copy.deepcopy(source_row._tr)
    table._tbl.append(new_tr)

    # Return as Row object
    from docx.table import _Row
    return _Row(new_tr, table)


def _capture_row_styles(ws: Worksheet, row: int) -> dict:
    """Capture cell styles from a row for replication."""
    styles = {}
    for cell in ws[row]:
        col = cell.column
        styles[col] = {
            "font": copy.copy(cell.font),
            "border": copy.copy(cell.border),
            "fill": copy.copy(cell.fill),
            "alignment": copy.copy(cell.alignment),
            "number_format": cell.number_format,
        }
    return styles


def _apply_row_styles(ws: Worksheet, row: int, styles: dict):
    """Apply captured styles to a row."""
    for col, style in styles.items():
        cell = ws.cell(row=row, column=col)
        cell.font = style["font"]
        cell.border = style["border"]
        cell.fill = style["fill"]
        cell.alignment = style["alignment"]
        cell.number_format = style["number_format"]
